"""
文档解析器

支持多种文档格式的解析：
- Word (.docx)
- PDF (.pdf)
- Markdown (.md)
- 文本 (.txt)
- HTML (.html)
- 云文档URL
"""

from pathlib import Path
from typing import Optional
import re

from src.core.config import Config
from src.core.exceptions import DocumentParseError
from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocumentParser:
    """
    文档解析器
    
    使用示例：
        ```python
        parser = DocumentParser(config)
        
        # 解析本地文件
        content = parser.parse_file("文档.docx")
        
        # 解析URL
        content = await parser.parse_url("https://docs.example.com/doc")
        ```
    """
    
    def __init__(self, config: Config):
        """
        初始化解析器
        
        Args:
            config: 配置对象
        """
        self.config = config.document
    
    def parse_file(self, file_path: str) -> str:
        """
        解析本地文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            文档文本内容
        
        Raises:
            DocumentParseError: 解析失败
        """
        path = Path(file_path)
        
        if not path.exists():
            raise DocumentParseError(f"文件不存在: {file_path}")
        
        # 检查文件大小
        file_size = path.stat().st_size
        if file_size > self.config.max_file_size:
            raise DocumentParseError(f"文件过大: {file_size} bytes")
        
        # 根据扩展名选择解析器
        suffix = path.suffix.lower()
        
        parsers = {
            ".docx": self._parse_docx,
            ".pdf": self._parse_pdf,
            ".md": self._parse_text,
            ".txt": self._parse_text,
            ".html": self._parse_html,
            ".htm": self._parse_html,
        }
        
        parser = parsers.get(suffix)
        if not parser:
            raise DocumentParseError(f"不支持的文件格式: {suffix}")
        
        try:
            content = parser(file_path)
            logger.info(f"Parsed file: {file_path}, content length: {len(content)}")
            return content
        except Exception as e:
            raise DocumentParseError(f"解析失败: {str(e)}")
    
    async def parse_url(self, url: str) -> str:
        """
        解析云文档URL
        
        支持：
        - 飞书文档
        - Confluence
        - 语雀
        - 其他可通过HTTP获取的文档
        
        Args:
            url: 文档URL
        
        Returns:
            文档文本内容
        """
        import httpx
        
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(url, timeout=30)
                response.raise_for_status()
                
                content_type = response.headers.get("content-type", "")
                
                # 如果是HTML页面，提取文本
                if "text/html" in content_type:
                    return self._extract_text_from_html(response.text)
                
                # 否则返回原始文本
                return response.text
                
        except Exception as e:
            raise DocumentParseError(f"获取URL失败: {str(e)}")
    
    def _parse_docx(self, file_path: str) -> str:
        """解析Word文档"""
        from docx import Document
        
        doc = Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        return "\n\n".join(paragraphs)
    
    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文档"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_parts = []
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return "\n\n".join(text_parts)
                
        except Exception as e:
            # 如果PyPDF2失败，尝试使用pdfplumber
            try:
                import pdfplumber
                
                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                
                return "\n\n".join(text_parts)
                
            except ImportError:
                raise DocumentParseError("请安装 PyPDF2 或 pdfplumber: pip install PyPDF2 pdfplumber")
    
    def _parse_text(self, file_path: str) -> str:
        """解析文本文件"""
        encodings = [self.config.encoding, "utf-8", "gbk", "gb2312", "latin-1"]
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise DocumentParseError(f"无法解码文件，尝试的编码: {encodings}")
    
    def _parse_html(self, file_path: str) -> str:
        """解析HTML文件"""
        with open(file_path, 'r', encoding=self.config.encoding, errors='ignore') as f:
            html = f.read()
        
        return self._extract_text_from_html(html)
    
    def _extract_text_from_html(self, html: str) -> str:
        """从HTML中提取文本"""
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(html, 'html.parser')
        
        # 移除脚本和样式
        for script in soup(["script", "style"]):
            script.decompose()
        
        # 获取文本
        text = soup.get_text()
        
        # 清理空白
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        
        return text
    
    def clean_text(self, text: str) -> str:
        """
        清理文本
        
        - 移除多余空白
        - 统一换行符
        - 移除特殊字符
        """
        # 统一换行符
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # 移除多余空白行
        lines = text.split('\n')
        cleaned_lines = []
        prev_empty = False
        
        for line in lines:
            stripped = line.strip()
            if stripped:
                cleaned_lines.append(stripped)
                prev_empty = False
            elif not prev_empty:
                cleaned_lines.append('')
                prev_empty = True
        
        return '\n'.join(cleaned_lines)
